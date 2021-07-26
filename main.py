import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
from matplotlib import pyplot as plt
import matplotlib.image as mpimg
import time
from docx2pdf import convert
from tkinter import Tk, filedialog


def select_folder():
    root = Tk()  # pointing root to Tk() to use it as Tk() in program.
    root.withdraw()  # Hides small tkinter window
    root.attributes('-topmost', True)  # Opened windows will be active. above all windows despite of selection
    open_file = filedialog.askdirectory()  # Returns opened path as str
    return open_file


# read csv
print("Please select id_list CSV file")
id_list_file = select_folder()
df = pd.read_csv(id_list_file)

while 1:
    # Check number of photo in folder
    from_dir = 'photo'
    from_arr = os.listdir(from_dir)
    if len(from_arr) < 8:
        print("All completed. Thank you Ching. Respect!!!")
        break

    # create figure
    fig = plt.figure(figsize=(8, 6))

    # setting values to rows and column variables
    rows = 4
    columns = 2

    # reading images
    Image1 = mpimg.imread(f'photo/{from_arr[0:1][0]}')
    Image2 = mpimg.imread(f'photo/{from_arr[1:2][0]}')
    Image3 = mpimg.imread(f'photo/{from_arr[2:3][0]}')
    Image4 = mpimg.imread(f'photo/{from_arr[3:4][0]}')
    Image5 = mpimg.imread(f'photo/{from_arr[4:5][0]}')
    Image6 = mpimg.imread(f'photo/{from_arr[5:6][0]}')
    Image7 = mpimg.imread(f'photo/{from_arr[6:7][0]}')
    Image8 = mpimg.imread(f'photo/{from_arr[7:8][0]}')
    Image9 = mpimg.imread(f'photo/{from_arr[8:9][0]}')

    # Adds a subplot at the 1st position
    fig.add_subplot(rows, columns, 1)

    # showing image
    plt.imshow(Image1)
    plt.axis('off')
    plt.title("First")

    # Adds a subplot at the 2nd position
    fig.add_subplot(rows, columns, 2)

    # showing image
    plt.imshow(Image2)
    plt.axis('off')
    plt.title("Second")

    # Adds a subplot at the 3rd position
    fig.add_subplot(rows, columns, 3)

    # showing image
    plt.imshow(Image3)
    plt.axis('off')
    plt.title("Third")

    # Adds a subplot at the 4th position
    fig.add_subplot(rows, columns, 4)

    # showing image
    plt.imshow(Image4)
    plt.axis('off')
    plt.title("Fourth")

    # Adds a subplot at the 5th position
    fig.add_subplot(rows, columns, 5)

    # showing image
    plt.imshow(Image5)
    plt.axis('off')
    plt.title("Fifth")

    # Adds a subplot at the 6th position
    fig.add_subplot(rows, columns, 6)

    # showing image
    plt.imshow(Image6)
    plt.axis('off')
    plt.title("Sixth")

    # Adds a subplot at the 7th position
    fig.add_subplot(rows, columns, 7)

    # showing image
    plt.imshow(Image7)
    plt.axis('off')
    plt.title("Seventh")

    # Adds a subplot at the 8th position
    fig.add_subplot(rows, columns, 8)

    # showing image
    plt.imshow(Image8)
    plt.axis('off')
    plt.title("Eighth")

    plt.ion()
    plt.show()
    plt.pause(0.001)

    skip_no = input("Input skip number:").upper()
    if len(skip_no) != 7:
        skip_no = input("Input correct skip number:").upper()

    plt.close()

    skip_company = df.loc[df['Number'] == skip_no]['Company'].iloc[0]

    document = Document()

    # paragraph style
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'Microsoft JhengHei'

    # heading style
    font_styles = document.styles
    font_charstyle = font_styles.add_style('HeadingStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(24)
    font_object.name = 'Microsoft JhengHei'
    heading = document.add_paragraph()
    heading.add_run('環保斗自願登記試驗計劃項目記綠', style='HeadingStyle')

    info1 = document.add_paragraph()
    info1.add_run(f'公司名稱：{skip_company}\n環保斗登記編號： {skip_no}', style='CommentsStyle')

    # add image to word
    for photo in from_arr[0:8]:
        document.add_picture(f"photo/{photo}", width=Inches(3))

    # create skip folder
    path = f"skip_word_folder/{skip_no}"
    os.mkdir(path)

    # save word to folder
    document.save(f'skip_word_folder/{skip_no}/{skip_no}.docx')

    print(f"{skip_no} {skip_company} word successfully created!!!")

    convert(f"skip_word_folder/{skip_no}/{skip_no}.docx", f"pdf/{skip_no}.pdf")
    time.sleep(2)
    print(f"\n{skip_no} {skip_company} pdf successfully created!!!\n")

    # del photo that are done
    for photo in from_arr[0:8]:
        os.remove(f"photo/{photo}")
